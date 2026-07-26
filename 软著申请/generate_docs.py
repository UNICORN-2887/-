"""生成软著源码文档：前30页+后30页，小五号Courier New字体"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT_DIR = r"E:\Project\DeadMaze"
OUTPUT_DIR = os.path.join(PROJECT_DIR, "软著申请")
SOFT_NAME = "DeadMaze\u6e38\u620f\u81ea\u52a8\u5316\u5bfc\u822a\u6218\u6597\u7cfb\u7edf"
VERSION = "V1.0"

SOURCE_FILES = [
    "navigator.py", "config_server.py", "calibrate.html",
    "map_stitcher.py", "reachability_map.py", "pathfinder.py",
    "map_tracker.py", "game_controller.py", "map_cropper.py",
    "map_localizer.py", "screenshots_capture.py", "threshold_viewer.py",
    "AImaneuver/combat_dashboard.py", "AImaneuver/ocr_reader.py",
    "AImaneuver/hp_detector.py", "AImaneuver/inventory_ocr.py",
    "AImaneuver/supply_check.py", "AImaneuver/supply_decision.py",
    "AImaneuver/supply_test_panel.py", "AImaneuver/camera_finder.py",
    "AImaneuver/campfire_calibrate.py", "AImaneuver/food_ocr_calibrate.py",
    "test/test_launcher.py",
]

def set_page_margins(doc, top=2.54, bottom=2.54, left=2.54, right=2.54):
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)

def add_header_footer(doc):
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(f"DeadMaze Automation {VERSION}")
        run.font.name = '\u5b8b\u4f53'
        run.font.size = Pt(7.5)

def generate_source_doc():
    doc = Document()
    set_page_margins(doc)
    add_header_footer(doc)

    all_lines = []
    for fname in SOURCE_FILES:
        fpath = os.path.join(PROJECT_DIR, fname)
        if not os.path.exists(fpath):
            print(f"[!] 跳过: {fname}")
            continue
        with open(fpath, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        all_lines.append(f"// ====== {fname} ======")
        in_docstring = False
        in_block_comment = False
        for line in content.split('\n'):
            stripped = line.strip()
            # 空行跳过
            if not stripped:
                continue
            # 文档字符串块 ("""...""")
            if stripped.startswith('"""') or stripped.startswith("'''"):
                if in_docstring:
                    in_docstring = False
                else:
                    in_docstring = True
                continue
            if in_docstring:
                continue
            # 块注释 /* ... */
            if stripped.startswith('/*'):
                in_block_comment = True
                continue
            if in_block_comment:
                if '*/' in stripped:
                    in_block_comment = False
                continue
            # 纯注释行 (Python #, JS //, HTML <!--)
            if (stripped.startswith('#') or stripped.startswith('//') or
                stripped.startswith('<!--')):
                continue
            if len(stripped) > 100:
                all_lines.append(stripped[:100])
            else:
                all_lines.append(stripped)

    total_lines = len(all_lines)
    lines_per_page = 50
    pages_needed = (total_lines + lines_per_page - 1) // lines_per_page
    print(f"总行数: {total_lines}, 每页{lines_per_page}行, 需要{pages_needed}页")

    front_lines = all_lines[:30 * lines_per_page]
    back_start = max(0, total_lines - 30 * lines_per_page)
    back_lines = all_lines[back_start:]

    # 前30页
    current_page = 1
    for i, line in enumerate(front_lines):
        if current_page > 30:
            break
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(10.5)
        run = p.add_run(f"{i+1:4d}  {line}")
        run.font.name = 'Courier New'
        run.font.size = Pt(6.5)
        if (i + 1) % lines_per_page == 0:
            doc.add_page_break()
            current_page += 1

    # 分隔页
    doc.add_page_break()
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep.add_run(f"（前30页结束，以下为后30页——源代码第{back_start+1}行至第{total_lines}行）")
    run.font.size = Pt(10)
    run.font.bold = True
    doc.add_page_break()

    # 后30页
    current_page = 1
    for i, line in enumerate(back_lines):
        if current_page > 30:
            break
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(10.5)
        actual_line = back_start + i + 1
        run = p.add_run(f"{actual_line:4d}  {line}")
        run.font.name = 'Courier New'
        run.font.size = Pt(6.5)
        if (i + 1) % lines_per_page == 0:
            doc.add_page_break()
            current_page += 1

    path = os.path.join(OUTPUT_DIR, f"DeadMaze Automation{VERSION}-\u6e90\u7801.doc")
    doc.save(path)
    print(f"\u6e90\u7801\u6587\u6863: {path}")
    return path

if __name__ == "__main__":
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    generate_source_doc()
    print("Done!")
