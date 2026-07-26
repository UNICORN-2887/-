"""生成软著用户操作手册 .doc"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT_DIR = r"E:\Project\DeadMaze"
OUTPUT_DIR = os.path.join(PROJECT_DIR, "软著申请")
SOFT_NAME = "DeadMaze 游戏自动化导航战斗系统"
VERSION = "V1.0"

def set_page_margins(doc, top=2.54, bottom=2.54, left=2.54, right=2.54):
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)

def generate_manual():
    doc = Document()
    set_page_margins(doc)

    # ===== 封面 =====
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(SOFT_NAME)
    run.font.size = Pt(22)
    run.font.bold = True
    doc.add_paragraph()
    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ver.add_run(f"用户操作手册 {VERSION}")
    run.font.size = Pt(16)
    doc.add_page_break()

    # ===== 正文 =====
    sections = [
        ("1. 系统概述", [
            f"{SOFT_NAME}是一款为DeadMaze 2.5D俯视角打僵尸游戏打造的综合性全自动辅助工具。系统包含光流法地图自动拼接、二值可达区标定编辑器、A*路径规划导航、YOLOv8僵尸实时检测、EasyOCR状态读取、智能补给决策引擎、后台SendMessage操控和网页配置面板等功能，实现游戏内寻路、战斗、补给、返航的全流程闭环自动化。",
            "本系统采用Python+Flask+HTML5架构，通过OBS虚拟摄像头采集游戏画面（1920x1080分辨率），利用OpenCV进行图像处理和定位追踪，YOLOv8进行僵尸目标检测，EasyOCR中英文双模型进行UI状态识别。系统通过pip一键安装依赖，bat批处理脚本启动，降低用户上手门槛。",
            "项目网站：https://blog.219882.xyz/deadmaze/",
            "开源仓库：https://github.com/UNICORN-2887/-",
        ]),
        ("2. 环境安装", [
            "2.1 Python 3.10+：访问 python.org 下载安装，安装时勾选「Add Python to PATH」。",
            "[配图：setup.bat运行界面截图，显示Python检测和pip安装依赖的过程]",
            "2.2 一键安装依赖：解压项目后双击 setup.bat，自动检测Python环境并pip安装opencv-python/numpy/pywin32/easyocr/ultralytics/flask/mss/pygrabber/psutil等依赖库",
            "2.3 OBS Studio：下载安装OBS Studio，添加游戏窗口为采集源，工具→虚拟摄像头→启动，确保输出分辨率为1920x1080。OBS虚拟摄像头是游戏画面采集的唯一方式。",
            "[配图：OBS Studio主界面，红框标出工具→虚拟摄像头→启动菜单路径，以及设置→视频→1920x1080分辨率]",
            "2.4 如果Python环境检测失败，请确认已安装Python 3.10+并勾选PATH选项；如果pip安装失败，可手动运行: pip install -r requirements.txt",
        ]),
        ("3. 快速开始", [
            "[配图：run_config.bat运行后的终端窗口，显示「Server ready」和浏览器自动打开的页面]",
            "3.1 配置与标定：双击 run_config.bat 启动网页配置服务，等待约6秒自动打开浏览器，同时打开配置面板(http://127.0.0.1:5050)和标定中心(/calibrate)两个页面",
            "3.2 设置游戏路径：在配置面板GAME区域填写DeadMaze游戏exe的完整路径",
            "3.3 摄像头选择：在标定中心顶部选择OBS虚拟摄像头，点击「测试」确认画面正常，点击「保存」",
            "3.4 ROI标定：在标定中心点击「实时预览」，右侧列表显示10个ROI区域(EXP/HUNGER/THIRST/STAMINA/THREAT/OPEN/HP/WEAPON/INVENTORY/FOOD)。点击卡片选中对应ROI→拖拽或WASD移动彩色框→检查识别结果→保存",
            "3.5 启动导航：双击 run_navigator.bat，选择地图(默认MazonAcademy)，在地图窗口左键设起点、右键设终点、Enter开始导航",
        ]),
        ("4. 导航使用", [
            "[配图：导航窗口全貌，包含地图、A*路径蓝色线条、起点/终点标记、状态面板]",
            "[配图：run_navigator.bat地图选择菜单，列出可用地图]",

            "4.1 设定起点：在地图窗口左键点击角色当前位置",
            "4.2 设定终点：右键点击目标位置，系统自动A*规划路径（蓝色线条），显示预估距离",
            "4.3 开始导航：按Enter键，角色自动沿规划路径8方向移动",
            "4.4 途径点：连续右键点击多个位置设置途径点，按M键开启循环巡逻，角色在各点间循环移动并在每点执行战斗",
            "4.5 返航：按H键手动返航，系统自动A*寻路回火堆，YOLO检测火堆位置后后台点击交互",
            "4.6 暂停/停止：空格键暂停/继续，Esc键停止当前导航，Q键退出程序",
            "4.7 技能：数字键1-4手动释放技能，E键切换技能自动释放开关",
            "4.8 地图操作：IJKL键平移地图视图，+/-键缩放地图，O/P键调节低状态返航阈值",
            "4.9 地图选择：双击run_navigator.bat后显示可用地图列表，输入数字选择，Enter确认",
        ]),
        ("5. 战斗系统", [
            "5.1 YOLO检测：系统每0.5秒通过YOLOv8模型检测画面中的僵尸目标，标记位置和置信度",
            "5.2 战斗规则优先级：全局补血(HP<80%→释放技能2)→HP过低脱战(HP<20%→空格脱战→返回途径点)→Threat≥2自动返航→武器空槽检测→低状态返航(H/T/S<阈值)→正常战斗/巡逻",
            "5.3 追击逻辑：僵尸数量≤MaxZombies且HP≥CombatEntryHP时进入战斗，在ZombieRange半径内搜索目标，攻击距离AttackRange，追击超时ChaseTimeout后换目标",
            "5.4 技能管理：技能栏第2格必须放置治疗技能，冷却时间在配置面板设置（默认4/12/22/32秒），系统按冷却自动释放技能1/3/4进行战斗",
            "5.5 武器管理：每15秒自动整理背包后检测武器第一格颜色（参考色RGB 80,39,19），颜色匹配率超过阈值判定为空槽→自动返航火堆→进入火堆后停止程序",
            "5.6 后台操控：所有键盘和鼠标操作通过Windows SendMessage API实现，不抢占前台焦点，不影响用户其他操作",
        ]),
        ("6. 补给系统", [
            "6.1 触发条件：饱食度/口渴度/耐力低于阈值（默认15）或威胁值≥2时自动触发返航补给",
            "6.2 返航流程：A*寻路回火堆→YOLO检测火堆位置→后台SendMessage点击→EasyOCR确认交互界面「开」字→进入补给决策",
            '6.3 背包扫描：自动拖拽8格食物栏，对每格OCR识别物品名称和数量，识别[食物]和[水]及其数值',
            "6.4 补给决策：计算最优食物组合使饱食度>100且口渴度>100，自动后台点击消耗食物",
            "6.5 循环补给：一轮不够自动继续，直至满足条件或无可食用物品，最多吃3个避免浪费",
            "6.6 补给后恢复：自动关闭交互界面，恢复之前的巡逻路线或导航任务",
        ]),
        ("7. 网页配置面板", [
            "[配图：网页配置面板全貌，展示NAVIGATION/COMBAT/STATUS/SKILLS/WEAPON各参数滑块]",

            "7.1 启动：双击run_config.bat或运行python config_server.py，浏览器访问http://127.0.0.1:5050",
            "7.2 导航参数：WP Reach/Deviation/Move Dur/Goal Reach/Lookahead共5个导航阈值，滑块调节实时生效",
            "7.3 战斗参数：Zombie Range/Attack Range/Chase Timeout/Combat Entry HP%/Max Zombies共5个战斗阈值",
            "7.4 状态参数：Low Stat Thr/Heal HP%/Escape HP%/Return Thr共4个状态触发阈值",
            "7.5 技能参数：Skill 1-4 CD共4个技能冷却时间（游戏内CD+2秒）",
            "7.6 武器参数：W Tolerance/W Threshold/W Check共3个武器空槽检测参数",
            "7.7 通知参数：PushPlus Token（在pushplus.plus获取），填入后停止/返航/低状态等事件微信推送通知",
            "7.8 启动器参数：游戏路径（DeadMaze.exe完整路径）、加速器路径（可选加速版插件路径）",
            "7.9 保存：点击SAVE CONFIG按钮保存所有参数到navigator_config.json",
        ]),
        ("8. ROI标定中心", [
            "[配图：标定中心全貌——左侧OBS画面+彩色ROI边框叠加，右侧三组ROI卡片列表，卡片下方显示OCR识别结果]",
            "[配图：摄像头选择下拉框和测试按钮，确认OBS画面正确]",

            "8.1 入口：启动config后访问http://127.0.0.1:5050/calibrate",
            "8.2 摄像头选择：页面顶部下拉框列出所有可用摄像头，OBS虚拟摄像头标注[OBS]，点击测试确认画面，点击保存",
            "8.3 实时预览：点击「▶实时预览」按钮，每2秒自动刷新OBS画面并运行OCR/HP/武器检测",
            "8.4 单张截图：点击「📷单张截图」按钮一次性截取当前画面进行分析",
            "8.5 ROI区域：右侧面板分三组列出10个标定项——OCR状态识别(EXP/HUNGER/THIRST/STAMINA/THREAT/OPEN)、战斗检测(HP/WEAPON)、背包补给(INVENTORY/FOOD)",
            "8.6 调整方式：点击右侧卡片选中ROI→图上对应彩色边框高亮→鼠标拖拽移动框→WASD/方向键微调(Shift+方向键加速)→手动输入X/Y/W/H数字微调",
            "8.7 识别验证：每个OCR卡片下方实时显示识别结果文字，对照游戏画面确认位置正确",
            "8.8 保存与重置：点击💾保存全部按钮写入JSON配置文件；点击↺重置按钮恢复为本地JSON文件值",
            "8.9 设备适配：OBS虚拟摄像头缩放/偏移因用户而异，标定页面使每个用户可独立调整ROI位置适配自己的OBS配置",
        ]),
        ("9. 自建地图", [
            "[配图：map_stitcher光流建图——左为OBS实时画面（带裁剪框），右为拼接中的大地图]",
            "[配图：reachability_map可达区标定——涂刷/描边/门标记操作界面，标注各模式和键位]",

            "9.1 光流建图：运行 python map_stitcher.py -c 1 -o MyMap.jpg",
            "9.2 裁剪框：按T显示/隐藏裁剪框，IJKL移动位置，+/-缩放边距，排除HUD和工具栏区域",
            "9.3 自动拼接：按A开启自动模式，操控角色沿地图边界走一圈，再蛇形填充内部区域，系统每0.3秒自动拼接一帧",
            "9.4 保存：按S保存地图和裁剪配置。输出为高分辨率全景图（约20000x10000像素）",
            "9.5 可达区标定：运行 python reachability_map.py MyMap.jpg -o MyMap_reachable.png",
            "9.6 涂刷模式（默认）：左键涂白（可达），右键涂黑（不可达），数字键1-4切换画笔大小（4/12/30/80像素）",
            "9.7 描边模式（P键）：点击绘制多边形顶点，Enter键填充，F键切换黑白填充颜色",
            "9.8 门标记模式（D键）：点击标记门的位置，1键设置左上↔右下方向，2键设置右上↔左下方向",
            "9.9 火堆标定：右键点击地图上火堆位置，自动保存坐标JSON",
            "9.10 HSV初稿：按C键自动用HSV颜色分割生成初稿，再手动修边提高效率",
            "9.11 导航验证：运行 python navigator.py MyMap_reachable.png --map MyMap.jpg，设起终点测试路径通畅性",
        ]),
        ("10. 地图社区", [
            "10.1 地图下载：访问网站 https://blog.219882.xyz/deadmaze/ → 地图下载章节，MazonAcademy和Lakeview18可直接下载",
            "10.2 地图目录：所有地图存放在项目map/子文件夹，每个地图一个文件夹（如map/MazonAcademy/）",
            "10.3 使用下载地图：解压ZIP到map/目录，确保形成map/地图名/文件夹结构",
            "10.4 提交地图：访问网站→提交地图页面→按指引打包文件夹为ZIP→命名格式「地图名_用户名_版本号.zip」→发送邮件到管理员邮箱",
            "10.5 邮件标题格式：[DeadMaze提交] 地图名 - 用户名 - 版本号",
            "10.6 审核流程：管理员运行admin_review.py自动扫描QQ邮箱→下载ZIP附件→验证文件结构→解压到地图库→发布到网站",
            "10.7 支持地图列表：MazonAcademy/Lakeview18/BodegaBay/BlueMesa/WalkerRiver/SunsetMall/SantaRosaDowntown/Highway99/ArizonaJurassicMuseum/SacramentoSuburbs/SurvivorCamp共11个地图",
        ]),
        ("11. 微信推送通知", [
            "11.1 配置：在pushplus.plus注册获取Token，填入网页配置面板NOTIFY区域并保存",
            "11.2 推送时机：武器耗尽→程序停止、HP过低→脱战、Threat≥2→自动返航、饱食/口渴/耐力低于阈值→返航补给",
            "11.3 限流机制：同一标题60秒内不重复发送，避免频繁推送",
            "11.4 推送内容包含HP/饱食/口渴/耐力/威胁等完整状态摘要",
        ]),
        ("12. 常见问题", [
            "问：启动后提示OBS摄像头未开？答：确认OBS Studio已运行，菜单→工具→虚拟摄像头→启动，输出分辨率1920x1080",
            "问：OCR识别不准？答：去标定页面/calibrate调整ROI位置，确认摄像头选择正确（OBS虚拟摄像头），不同OBS缩放/偏移需要重新标定",
            "问：导航时角色乱走/卡住？答：检查可达图标定是否完整，偏离路径会自动重规划；如果频繁重规划说明可达区标定有问题",
            "问：HP一直0%？答：去标定页面调整HP ROI框位置，确保框内是绿色血量条区域",
            "问：武器检测不准？答：调整武器ROI框到武器第一格正上方，可在标定页面微调Tol(容差)和Thr(阈值)参数",
            "问：如何在其他地图使用？答：按照第9章自建地图流程操作：光流建图→可达区标定→火堆标记→导航验证",
            "问：技能2不自动释放？答：确认技能栏第2格放置的是治疗技能，Heal HP%参数设置合理（默认80%）",
            "问：config页面打开后无法连接？答：Flask服务需要约6秒启动，等待片刻刷新页面即可；run_config.bat会自动等待6秒后打开浏览器",
            "问：标定页面实时预览画面不动？答：确认选中OBS虚拟摄像头而非内置摄像头；独立模式使用持久摄像头连接",
            "问：建图时地图拼接不完整？答：确保绕边界完整走一圈，室内区域多走几遍增加重叠帧，检查裁剪框是否排除了游戏画面",
        ]),
    ]

    for title_text, paragraphs in sections:
        doc.add_heading(title_text, level=1)
        for p_text in paragraphs:
            para = doc.add_paragraph(p_text)
            for run in para.runs:
                run.font.name = '\u5b8b\u4f53'
                run.font.size = Pt(12)
        doc.add_paragraph()

    doc.add_paragraph()
    doc.add_heading("技术支持", level=1)
    doc.add_paragraph("项目网站：https://blog.219882.xyz/deadmaze/")
    doc.add_paragraph("开源仓库：https://github.com/UNICORN-2887/-")

    # ===== 配图清单 =====
    # 请在Word中搜索 [配图： 找到所有占位，手动插入对应截图后删除占位文字
    #
    # 图片1: screen-config.png — 网页配置面板
    #   → 第7章 网页配置面板
    # 图片2: screen-calibrate.png — ROI标定中心
    #   → 第8章 ROI标定中心
    # 图片3: screen-navigator.png — 导航窗口
    #   → 第4章 导航使用
    # 图片4: screen-stitching.png — 光流法建图
    #   → 第9章 自建地图
    # 图片5: screen-reachable.png — 可达区标定
    #   → 第9章 自建地图
    # 图片6: setup.bat运行截图
    #   → 第2章 环境安装
    # 图片7: run_config.bat终端窗口
    #   → 第3章 快速开始
    # 图片8: OBS Studio设置界面
    #   → 第2章 环境安装
    # 图片9: 地图选择菜单
    #   → 第4章 导航使用
    # 图片10: 摄像头选择器
    #   → 第8章 ROI标定中心

    path = os.path.join(OUTPUT_DIR, f"DeadMaze Automation\u7528\u6237\u64cd\u4f5c\u624b\u518c-{VERSION}.doc")
    doc.save(path)
    print(f"用户手册: {path}")
    return path

if __name__ == "__main__":
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    generate_manual()
    print("Done!")
